"""
Word (.docx) export for generated documents (research briefs, risk
assessments, workflow write-ups, ...).

`generated_docs`/`agent_catalog` only ever produce body_html restricted
to a strict tag allowlist (see agent_catalog.sanitize_html's
_ALLOWED_TAGS: ul, ol, li, b, i, em, strong, br, p, div, span) — small
enough to walk directly with the stdlib HTMLParser instead of pulling in
a general-purpose HTML-to-docx conversion library.
"""

from __future__ import annotations

import io
import re
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt

_BLOCK_TAGS = {'p', 'div', 'li'}
_BOLD_TAGS = {'b', 'strong'}
_ITALIC_TAGS = {'i', 'em'}


class _BodyHtmlToDocx(HTMLParser):
    """Walks the sanitized body_html and writes paragraphs/runs straight
    into a python-docx Document. list_depth tracks ul/ol nesting so
    nested bullets still read as nested (indented), not flattened."""

    def __init__(self, doc: Document):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.bold = 0
        self.italic = 0
        self.list_stack: list[str] = []   # 'ul' | 'ol'
        self.list_counts: list[int] = []
        self.para = None

    def _ensure_para(self, style: str | None = None):
        if self.para is None:
            self.para = self.doc.add_paragraph(style=style)
        return self.para

    def handle_starttag(self, tag, attrs):
        if tag in _BOLD_TAGS:
            self.bold += 1
        elif tag in _ITALIC_TAGS:
            self.italic += 1
        elif tag == 'br':
            self._ensure_para().add_run().add_break()
        elif tag in ('ul', 'ol'):
            self.list_stack.append(tag)
            self.list_counts.append(0)
        elif tag == 'li':
            style = 'List Bullet' if (self.list_stack and self.list_stack[-1] == 'ul') else 'List Number'
            depth = len(self.list_stack)
            style = f'{style} {depth}' if depth > 1 and f'{style} {depth}' in self.doc.styles else style
            self.para = self.doc.add_paragraph(style=style if style in self.doc.styles else None)
        elif tag in _BLOCK_TAGS:
            self.para = None
            self._ensure_para()

    def handle_endtag(self, tag):
        if tag in _BOLD_TAGS:
            self.bold = max(0, self.bold - 1)
        elif tag in _ITALIC_TAGS:
            self.italic = max(0, self.italic - 1)
        elif tag in ('ul', 'ol'):
            if self.list_stack:
                self.list_stack.pop()
                self.list_counts.pop()
        if tag in _BLOCK_TAGS or tag == 'li':
            self.para = None

    def handle_data(self, data):
        text = re.sub(r'\s+', ' ', data)
        if not text.strip():
            return
        run = self._ensure_para().add_run(text)
        run.bold = self.bold > 0
        run.italic = self.italic > 0


def html_to_docx_bytes(title: str, body_html: str, *, meta: list[tuple[str, str]] | None = None) -> bytes:
    """Render one generated document as a .docx file, returned as raw
    bytes ready to stream in an HTTP response. `meta` is an optional list
    of (label, value) lines placed under the title (author, category,
    confidence, date) — purely descriptive, never affects the body."""
    doc = Document()
    doc.add_heading(title or 'Document', level=1)
    if meta:
        p = doc.add_paragraph()
        run = p.add_run(' · '.join(f'{label}: {value}' for label, value in meta if value))
        run.italic = True
        run.font.size = Pt(9)
    parser = _BodyHtmlToDocx(doc)
    parser.feed(body_html or '')
    parser.close()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
